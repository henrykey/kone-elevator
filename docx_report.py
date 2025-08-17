import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches

def merge_json_to_docx(json_path, docx_template_path, output_path):
    # Load JSON data
    with open(json_path, 'r', encoding='utf-8') as f:
        json_data = json.load(f)
    
    test_results = json_data.get('test_results', [])
    if not test_results:
        raise ValueError("No 'test_results' found in JSON file.")
    
    # Create a dictionary for quick lookup: test number -> test details
    results_dict = {}
    for result in test_results:
        test_key = result.get('test', '').strip()
        if test_key.startswith('Test '):
            test_num_str = test_key[5:].strip()
            if test_num_str.isdigit():
                test_num = int(test_num_str)
                results_dict[test_num] = {
                    'test_result': result.get('test_result', 'N/A').strip(),
                    'status': result.get('status', 'N/A').strip(),
                    'description': result.get('description', 'N/A').strip(),
                    'duration_ms': result.get('duration_ms', 'N/A'),
                    'error_message': result.get('error_message', None),
                    'response_data': result.get('response_data', {}),
                    'api_calls': result.get('api_calls', []),
                    'request_parameters': result.get('request_parameters', {}),
                    'request_timestamp': result.get('request_timestamp', 'N/A'),
                    'response_timestamp': result.get('response_timestamp', 'N/A')
                }
    
    # Load the DOCX template
    doc = Document(docx_template_path)
    
    # Iterate through all tables in the document
    for table in doc.tables:
        test_num = None
        result_cell = None
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.replace('\n', ' ').strip()
            if label.startswith('Test ') and label[5:].isdigit():
                test_num = int(label[5:])
            elif label == 'Test result':
                result_cell = row.cells[1]
            if test_num is not None and result_cell is not None:
                # Get test details
                test_details = results_dict.get(test_num, {})
                
                # Build formatted result text with consistent indentation
                result_text = []
                result_text.append(f"Result: {test_details.get('test_result', 'N/A')}")
                result_text.append(f"Status: {test_details.get('status', 'N/A')}")
                result_text.append(f"Description: {test_details.get('description', 'N/A')}")
                result_text.append(f"Duration: {test_details.get('duration_ms', 'N/A')} ms")
                
                if test_details.get('error_message'):
                    result_text.append(f"Error Message: {test_details.get('error_message')}")
                
                # Add response data
                response_data = test_details.get('response_data', {})
                if response_data:
                    result_text.append("Response Data:")
                    result_text.append(json.dumps(response_data, indent=4, ensure_ascii=False))
                
                # Add API call details
                api_calls = test_details.get('api_calls', [])
                if api_calls:
                    result_text.append("API Calls:")
                    for i, call in enumerate(api_calls, 1):
                        result_text.append(f"  Call {i}:")
                        result_text.append(f"    Interface Type: {call.get('interface_type', 'N/A')}")
                        result_text.append(f"    URL: {call.get('url', 'N/A')}")
                        result_text.append(f"    Method: {call.get('method', 'N/A')}")
                        result_text.append(f"    Request Parameters:")
                        result_text.append(json.dumps(call.get('request_parameters', {}), indent=6, ensure_ascii=False))
                        result_text.append(f"    Response Data:")
                        result_text.append(json.dumps(call.get('response_data', {}), indent=6, ensure_ascii=False))
                        result_text.append(f"    Status Code: {call.get('status_code', 'N/A')}")
                        result_text.append(f"    Timestamp: {call.get('timestamp', 'N/A')}")
                        if call.get('error_message'):
                            result_text.append(f"    Error Message: {call.get('error_message')}")
                
                # Add request parameters
                request_params = test_details.get('request_parameters', {})
                if request_params:
                    result_text.append("Request Parameters:")
                    result_text.append(json.dumps(request_params, indent=4, ensure_ascii=False))
                
                result_text.append(f"Request Timestamp: {test_details.get('request_timestamp', 'N/A')}")
                result_text.append(f"Response Timestamp: {test_details.get('response_timestamp', 'N/A')}")
                
                # Clear existing content in the result cell
                for paragraph in result_cell.paragraphs:
                    paragraph.clear()
                
                # Add new paragraphs with formatted text, setting font size to 9 and line spacing to 1.0
                for line in result_text:
                    p = result_cell.add_paragraph(line.strip())
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    p.paragraph_format.line_spacing = 1.0  # Set line spacing to 1.0
                    for run in p.runs:
                        run.font.size = Pt(9)  # Set font size to 9 for inserted text only
                
                # Reset for next potential test
                test_num = None
                result_cell = None
    
    # Save the updated document
    doc.save(output_path)
    print(f"Updated DOCX saved to: {output_path}")

# Usage example (replace with actual file paths)
if __name__ == "__main__":
    json_path = "reports/validation_report.json"
    docx_template_path = "Service Robot API solution validation test guide version 2.docx"
    output_path = "formal_validation_report.docx"
    merge_json_to_docx(json_path, docx_template_path, output_path)